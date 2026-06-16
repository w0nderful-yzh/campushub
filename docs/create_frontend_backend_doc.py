from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "docs" / "chapter8-assets"
OUTPUT = ROOT / "docs" / "前端页面及对应前后端代码.docx"

CN_BODY = "宋体"
CN_HEADING = "黑体"
CODE_FONT = "Consolas"
BLACK = RGBColor(0, 0, 0)
MUTED = RGBColor(90, 90, 90)


def font(run, cn=CN_BODY, latin="Times New Roman", size=11, bold=False, color=BLACK):
    run.font.name = latin
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), cn)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)


def shade(cell, fill):
    props = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    props.append(shd)


def cell_margin(cell, value=100):
    props = cell._tc.get_or_add_tcPr()
    margins = OxmlElement("w:tcMar")
    for side in ("top", "start", "bottom", "end"):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        margins.append(node)
    props.append(margins)


def set_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for index, width in enumerate(widths):
            cell = row.cells[index]
            cell.width = Inches(width)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(round(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_heading(doc, text, level):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    font(run, cn=CN_HEADING, latin="Arial", size={1: 17, 2: 14}[level], bold=True)


def add_text(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Pt(22)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    font(p.add_run(text), size=11)


def add_picture(doc, filename, caption, width):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    p.add_run().add_picture(str(ASSETS / filename), width=Inches(width))
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(8)
    font(cp.add_run(caption), size=10)


def add_code(doc, title, code, page_break=False):
    if page_break:
        doc.add_page_break()
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.keep_with_next = True
    cp.paragraph_format.space_before = Pt(3)
    cp.paragraph_format.space_after = Pt(3)
    font(cp.add_run(title), size=10)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(6)
    p.paragraph_format.right_indent = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    ppr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F5F7FA")
    ppr.append(shd)
    borders = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "5")
        border.set(qn("w:color"), "D9DEE7")
        borders.append(border)
    ppr.append(borders)

    for i, line in enumerate(code.strip().splitlines()):
        if i:
            p.add_run().add_break()
        run = p.add_run(line or " ")
        font(run, cn=CN_BODY, latin=CODE_FONT, size=7.2, color=RGBColor(31, 41, 55))


def add_environment_table(doc):
    rows = [
        ("层次", "技术与运行环境"),
        ("前端", "Vue 3、TypeScript、Vite、Naive UI、Pinia、Vue Router、Axios"),
        ("后端", "Spring Boot、MyBatis-Plus、JWT、Maven、JDK 17"),
        ("数据服务", "MySQL 8.0、Redis 7、Elasticsearch 8.13.4"),
        ("通信方式", "前后端分离，通过 RESTful API 和 JSON 数据交互"),
        ("开发地址", "前端 http://localhost:5173；后端 http://localhost:8080"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    set_widths(table, [1.3, 5.0])
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            cell = table.cell(r, c)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_margin(cell)
            if r == 0:
                shade(cell, "E8EEF5")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c == 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            font(run, size=10, bold=r == 0)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CN_BODY)

    for level, size in ((1, 17), (2, 14)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLACK
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CN_HEADING)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(header.add_run("CampusHub 前后端实现"), size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(footer.add_run("CampusHub 校园社区系统"), size=9, color=MUTED)

    add_heading(doc, "8 前端页面及对应前后端代码", 1)
    add_heading(doc, "8.1 实验环境说明", 2)
    add_text(
        doc,
        "本实验采用前后端分离架构。前端负责页面展示、表单校验和用户交互，后端负责业务处理、身份认证和数据库访问。前后端通过 RESTful API 传输 JSON 数据，具体环境如下。",
    )
    add_environment_table(doc)

    add_heading(doc, "8.2 登录功能", 2)
    add_picture(doc, "login-interface.png", "图 8-1  登录界面", 4.7)
    add_text(doc, "用户填写用户名和密码后，前端完成非空校验并调用登录接口。后端查询用户、校验密码，成功后生成 JWT 令牌并返回用户信息。")

    add_code(
        doc,
        "代码 8-1  登录功能前端代码",
        """const rules: FormRules = {
  username: [{ required: true, message: '请输入用户名' }],
  password: [{ required: true, message: '请输入密码' }]
}

async function handleSubmit() {
  try {
    await formRef.value?.validate()
    submitting.value = true
    await authStore.loginAction({
      username: form.username.trim(),
      password: form.password
    })
    message.success('登录成功')
    router.replace('/')
  } finally {
    submitting.value = false
  }
}

export function loginApi(data: LoginDTO) {
  return request.post('/auth/login', data)
}""",
        page_break=True,
    )

    add_code(
        doc,
        "代码 8-2  登录功能后端代码",
        """@PostMapping("/login")
public ResponseEntity<Result> login(
        @RequestBody LoginDTO loginDTO) {
    Result result = authService.login(loginDTO);
    return ResponseEntity.status(result.getCode()).body(result);
}

public Result login(LoginDTO dto) {
    User user = userMapper.selectOne(
        new LambdaQueryWrapper<User>()
            .eq(User::getUsername, dto.getUsername()));
    if (user == null) {
        return Result.fail(404, "user not found");
    }
    if (!passwordEncoder.matches(
            dto.getPassword(), user.getPassword())) {
        return Result.fail(401, "login failed");
    }
    String token = JwtUtil.generateToken(
        user.getId(), user.getUsername());
    stringRedisTemplate.opsForValue().set(
        "login:token:" + user.getId(),
        token, 7, TimeUnit.DAYS);
    UserInfoVO userInfo = new UserInfoVO();
    BeanUtils.copyProperties(user, userInfo);
    userInfo.setRole(String.valueOf(user.getRole()));

    LoginVO loginVO = new LoginVO();
    loginVO.setToken(token);
    loginVO.setUserInfo(userInfo);
    return Result.ok(loginVO);
}""",
    )

    add_heading(doc, "8.3 帖子数据管理功能", 2)
    add_picture(doc, "post-management-interface.png", "图 8-2  帖子发布与管理界面", 4.65)
    add_text(
        doc,
        "帖子管理模块支持新增、查询、修改和删除。发布与编辑共用同一个表单页面，根据路由中是否存在帖子编号判断操作类型；删除操作在帖子详情页完成。",
    )

    add_code(
        doc,
        "代码 8-3  帖子管理前端代码",
        """const isEdit = computed(() => Boolean(route.params.id))

async function handleSubmit() {
  await formRef.value?.validate()
  const data = {
    categoryId: Number(form.categoryId),
    title: form.title.trim(),
    content: form.content.trim(),
    images: form.images.map(v => v.trim()).filter(Boolean)
  }
  if (isEdit.value) {
    await updatePostApi(editPostId.value, data)
  } else {
    await createPostApi(data)
  }
}

export const createPostApi = (data: SavePostDTO) =>
  request.post('/posts', data)
export const updatePostApi = (id: number, data: SavePostDTO) =>
  request.put(`/posts/${id}`, data)
export const deletePostApi = (id: number) =>
  request.delete(`/posts/${id}`)""",
    )

    add_code(
        doc,
        "代码 8-4  帖子接口控制层代码",
        """@PostMapping
public Result createPost(@RequestBody CreatePostDTO dto) {
    postService.createPost(dto);
    return Result.ok();
}

@PutMapping("/{postId}")
public Result updatePost(
        @PathVariable Long postId,
        @RequestBody UpdatePostDTO dto) {
    return postService.updatePost(postId, dto);
}

@DeleteMapping("/{postId}")
public Result deletePost(@PathVariable Long postId) {
    return postService.deletePost(postId);
}""",
    )

    add_code(
        doc,
        "代码 8-5  帖子新增服务代码",
        """@Transactional(rollbackFor = Exception.class)
public void createPost(CreatePostDTO dto) {
    Long userId = UserContext.getUserId();
    Post post = new Post()
        .setUserId(userId)
        .setCategoryId(dto.getCategoryId())
        .setTitle(dto.getTitle())
        .setContent(dto.getContent())
        .setViewCount(0)
        .setLikeCount(0)
        .setCommentCount(0)
        .setFavoriteCount(0)
        .setIsDeleted(0)
        .setCreateTime(LocalDateTime.now());
    save(post);
    // 将图片地址按顺序写入 image 表
    eventPublisher.publishEvent(
        new PostSearchSyncEvent(
            post.getId(), Action.UPSERT));
}""",
    )

    add_code(
        doc,
        "代码 8-6  帖子修改与删除服务代码",
        """public Result updatePost(Long postId, UpdatePostDTO dto) {
    Long userId = UserContext.getUserId();
    Post post = getById(postId);
    if (post == null || post.getIsDeleted() == 1)
        return Result.fail("帖子不存在");
    if (!post.getUserId().equals(userId))
        return Result.fail("无权修改此帖子");

    post.setCategoryId(dto.getCategoryId());
    post.setTitle(dto.getTitle());
    post.setContent(dto.getContent());
    updateById(post);
    // 删除旧图片并写入新图片
    return Result.ok();
}

public Result deletePost(Long postId) {
    Long userId = UserContext.getUserId();
    Post post = getById(postId);
    if (post == null || !post.getUserId().equals(userId))
        return Result.fail("无权删除此帖子");
    post.setIsDeleted(1);
    updateById(post);
    eventPublisher.publishEvent(
        new PostSearchSyncEvent(postId, Action.DELETE));
    return Result.ok();
}""",
    )

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
